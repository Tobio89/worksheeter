import docx, random, os
from docx.shared import Cm
# from .CNN import getNewsArticles, getArticleContent
from .get_news import getArticleContent
# from .vocab import getRandomUniqueWords, getMultipleDefinitions
from .comprehension import produceCloze
from .worksheet_config import discussion_questions, comprehension_questions



# MAIN_PATH = './app/static/download'
MAIN_PATH = os.path.join('.', 'app', 'static', 'download')

questions = []

def capitalizeEveryWord(s):

    s_split = s.split()

    s_capitalize = [w.capitalize() for w in s_split]

    return ' '.join(s_capitalize)

def addTitlePage(title, doc):
    
    doc.add_picture(os.path.join('.', 'app', 'static', 'ESL-auto_header.png'), width=Cm(14))
    
    doc.add_paragraph()
    doc.add_paragraph()

    doc.add_heading(capitalizeEveryWord(title), 0)

    return doc

def addVocaChunk(wordDict, doc):

    
    for word in wordDict:
        doc.add_heading(word, 3)
        for POS in wordDict[word]:

            if POS != 'examples':
            
                doc.add_heading(POS, 5)
                for meaning in wordDict[word][POS]:
                    if meaning[0].islower():
                        meaning = f'{meaning[0].upper()}{meaning[1:]}'

                    p = doc.add_paragraph()
                    p.style = 'ListBullet'
                    p.paragraph_format.left_indent = Cm(2)
                    p.text = meaning
                    


    return doc

def separateExamples(wordDict):

    # Produce list of examples, cloze-ified
    examples = []
    answer_words = []
    for word in wordDict:
        if wordDict[word]['examples']:
            answer_words.append(word)
            examples.append(wordDict[word]['examples'])

    
    random.shuffle(examples)
    random.shuffle(answer_words)

            
    
    return examples, answer_words


def addVocabClozeQuestions(wordDict, doc):

    question_source, answer_words = separateExamples(wordDict)

    random.shuffle(question_source)

    doc.add_heading('Vocab Questions',3 )
    doc.add_paragraph('')
    doc.add_heading('Fill in the blanks using the words in the table', 5)
    doc.add_paragraph('')
    
    words_table = doc.add_table(rows=1, cols=len(answer_words))
    words_table.style = 'TableGrid'
    word_row = words_table.rows[0].cells

    for i, word in enumerate(answer_words):
        word_row[i].text = word
    doc.add_paragraph('')
    for number, question in enumerate(question_source):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.text = f'{number+1}: {question}'

    return doc

def addDiscussionQuestions(listOfQuestions, doc):
    if not listOfQuestions:
        listOfQuestions = discussion_questions
    doc.add_heading('Discussion Questions',3 )
    doc.add_paragraph('')
    for number, question in enumerate(listOfQuestions):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.text = f'{number+1}: {question}'

    return doc


def addParagraphClozeQuestions(para_list, doc):
    questions = produceCloze(para_list)

    doc.add_heading('Reading Comprehension', 3)
    doc.add_paragraph('')
    doc.add_heading('Fill in the blanks', 5)
    doc.add_paragraph('')
    for question in questions:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.text = f'-: {question}'
        doc.add_paragraph()       
    
    return doc

def addReadingCompTemplateQuestions(doc):

    doc.add_heading('Comprehension Questions', 5)
    doc.add_paragraph()
    for question in comprehension_questions:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.text = f'{question}'
        doc.add_paragraph()
    
    return doc


def writeDocx(list_format_article, definitionDict, title):

    # Produce a title with sensible length
    alnumTitle = ''.join([i for i in title if i.isalnum() or i==' ']) # Remove special characters that prevent file saving.
    if len(alnumTitle.split(' ')) > 12:
        filename_title = ' '.join(alnumTitle.split(' ')[:12])
    else:
        filename_title = alnumTitle

    filename_DOCX = f'{filename_title}_Discussion_Worksheet.docx'

    # Initialise doc and write article to doc.
    doc = docx.Document()

    # Add title page
    doc = addTitlePage(title, doc)
    doc.add_page_break()

    # Add article section
    doc.add_heading('Article Reading', 1)
    doc.add_paragraph('')
    doc.add_heading(title, 2)
    doc.add_paragraph('')
    for para in list_format_article:
        doc.add_paragraph(para)
    doc.add_page_break()


    # Vocabulary Section

    doc.add_heading('Vocabulary Section', 1)

    # Add vocab glossary
    doc = addVocaChunk(definitionDict, doc) #Add main vocab definitions part
    doc.add_page_break()
    
    doc.add_heading('Comprehension Questions', 1)
    doc.add_paragraph('')
    doc.add_heading('Vocabulary Questions', 2)
    doc.add_paragraph('')
    # Add vocab cloze questions
    doc = addVocabClozeQuestions(definitionDict, doc)
    doc.add_page_break()

    doc.add_heading('Article Questions', 2)
    doc.add_paragraph('')
    # Add paragraph cloze questions
    doc = addParagraphClozeQuestions(list_format_article, doc)
    doc.add_page_break()

    # Add reading comp template questions
    doc = addReadingCompTemplateQuestions(doc)
    doc.add_page_break()

    doc.add_heading('Discussion Questions', 1)
    doc.add_paragraph('')
    # Add discussion questions
    doc = addDiscussionQuestions(questions, doc)

    print('writing document...')
    print(filename_DOCX)
    output_path = os.path.join(MAIN_PATH, filename_DOCX)
    doc.save(output_path)

    return filename_DOCX

if __name__ == "__main__":

    testSearchTerm = 'manchester united'
    try:
        news_arts = getNewsArticles(testSearchTerm)
    except:
        print('Bad search term :(')
        exit()

    while True:
        test_article = random.choice(news_arts)

        articleTitle = test_article['head']
        article_URL = test_article['href']

        print(f'Picked article: {articleTitle}')

        paragraph_divided = getArticleContent(article_URL)
        try:
            vocab_words = getRandomUniqueWords(paragraph_divided)
        except:
            continue
        paragraph_divided = getArticleContent(article_URL)

        writeDocx(paragraph_divided, vocab_words, articleTitle)

        break


